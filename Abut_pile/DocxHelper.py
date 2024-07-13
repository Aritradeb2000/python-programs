
import math
import cmath

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor,Pt, Inches
import os
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_UNDERLINE
 
document = Document()

def save_document(output_path):
    """Save the current state of the document to disk."""
    document.save(output_path)

def load_document(output_path):
    """Load the document from disk into memory."""
    global document
    document = Document(output_path)
    return document


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



def add_centered_image(document, image_path, width, height=None):
    """
    Add an image to a document with center alignment and specified width and height.
    
    Parameters:
        document (docx.Document): The document to which the image will be added.
        image_path (str): The file path of the image.
        width (float): The desired width of the image in inches.
        height (float, optional): The desired height of the image in inches. If not provided, 
            the height will be adjusted to maintain the aspect ratio.
    """
    # Add a paragraph
    p = document.add_paragraph()
    
    # Add a run to the paragraph and insert the picture
    run = p.add_run()
    run.add_picture(image_path)

    # Get the current width and height of the image
    image_width = run.element.xpath('.//wp:extent')[0].get('cx')
    image_height = run.element.xpath('.//wp:extent')[0].get('cy')
    image_width = int(image_width)
    image_height = int(image_height)

    # Calculate the current and desired width in EMU (English Metric Unit)
    current_width_emu = image_width
    desired_width_emu = int(width * 914400)  # Convert inches to EMU

    # If height is not provided, adjust it to maintain the aspect ratio
    if height is None:
        scaling_factor = desired_width_emu / current_width_emu
        desired_height_emu = int(image_height * scaling_factor)
    else:
        # Calculate the desired height in EMU
        desired_height_emu = int(height * 914400)  # Convert inches to EMU
        scaling_factor = min(desired_width_emu / current_width_emu, desired_height_emu / image_height)

    # Set the width and height of the image
    run.element.xpath('.//wp:extent')[0].set('cx', str(int(image_width * scaling_factor)))
    run.element.xpath('.//wp:extent')[0].set('cy', str(int(image_height * scaling_factor)))

    # Center align the paragraph
    p.alignment = 1  # 1 corresponds to center alignment in python-docx


def add_main_heading(document, heading_text):
    """
    Add a heading to the given document.

    Parameters:
        document (docx.Document): The Document object to which the heading will be added.
        heading_text (str): The text of the heading.
    """
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align the paragraph to center
    heading = p.add_run(heading_text)
    heading.bold = True
    heading.font.size = Pt(16)
    
    
    

def add_sub_heading(document, heading_text):
    """
    Add a heading to the given document.

    Parameters:
        document (docx.Document): The Document object to which the heading will be added.
        heading_text (str): The text of the heading.
    """
    p = document.add_paragraph()
    heading = p.add_run(heading_text)
    heading.bold = True
    heading.font.size = Pt(14)
    
    


def add_minor_heading(document, heading_text):
    """
    Add a heading to the given document.

    Parameters:
        document (docx.Document): The Document object to which the heading will be added.
        heading_text (str): The text of the heading.
    """
    p = document.add_paragraph()
    heading = p.add_run(heading_text)
    heading.bold = True
    heading.font.size = Pt(12)
   

def add_right_align_text(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_green_bold_text(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run("This text is bold and in deep green color.")
    run.bold = True
    font = run.font
    font.color.rgb = RGBColor(0, 128, 0)

def add_red_bold_text(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run("This text is bold and in deep green color.")
    run.bold = True
    font = run.font
    font.color.rgb = RGBColor(128, 0, 0)





def add_table_from_data(doc, data):
    # Add a table with 1 row and 3 columns
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    table.columns[0].width = 4000000  # Adjust width of the first column to be most wide
    table.columns[1].width = 2000000
    table.columns[2].width = 800000  # Adjust width of the last column to be less wide

    # Add table headers
    headers = table.rows[0].cells
    headers[0].text = "Description"
    headers[1].text = "Value"
   

    # Set header alignment and font size
    for cell in headers:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)

    # Add data rows
    for item in data:
        row = table.add_row().cells
        for i, value in enumerate(item):
            row[i].text = str(value)

            # Adjust spacing between paragraphs to decrease gap between rows
            for paragraph in row[i].paragraphs:
                paragraph.space_after = Pt(0)

        # Set row height
        table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[-1].height = Pt(14)  # Set row height to 10 points

    # Align text in cells and set font size
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
def add_table_to_doc(document, data, headers=None, index=None):
    """
    Add a table to a Word document.

    Parameters:
        document (docx.Document): The Word document object.
        data (list of lists): The 2D array of data for the table.
        headers (list): Optional. The headers for the table.
        index (list): Optional. The index for the table.

    Returns:
        None
    """
    # Create a table with one more row and one more column for headers and index
    num_rows = len(data) + (1 if headers else 0)
    num_cols = len(data[0]) + (1 if index else 0)
    table = document.add_table(rows=num_rows, cols=num_cols)

    # Set the style of the table to "Table Grid"
    table.style = 'Table Grid'

    # Add headers to the first row if provided
    if headers:
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align header text to center

    # Add index to the first column if provided
    if index:
        for i, index_value in enumerate(index):
            table.cell(i + (1 if headers else 0), 0).text = str(index_value)
            table.cell(i + (1 if headers else 0), 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align index text to center

    # Populate the rest of the table with data
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            cell = table.cell(i + (1 if headers else 0), j + (1 if index else 0))
            if isinstance(value, float):
                if(abs(value)>1):
                    value = round(value,2)
                elif(value==0):
                    value = round(value, 0)
                else: 
                   value = round(value, 3)

            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell content to center

    # Adjust the width of columns
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)


def add_table_to_docv2(document, data, headers=None, index=None):
    """
    Add a table to a Word document.

    Parameters:
        document (docx.Document): The Word document object.
        data (list of lists): The 2D array of data for the table.
        headers (list): Optional. The headers for the table.
        index (list): Optional. The index for the table.

    Returns:
        None
    """
    # Create a table with one more row and one more column for headers and index
    num_rows = len(data) + (1 if index else 0)
    num_cols = len(data[0]) + (1 if headers else 0)
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Add headers to the first row if provided
    if headers:
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i + (1 if index else 0)].text = header

    # Add index to the first column if provided
    if index:
        for i, index_value in enumerate(index):
            table.cell(i + 1, 0).text = str(index_value)

    # Populate the rest of the table with data
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            cell = table.cell(i + (1 if index else 0), j + (1 if headers else 0))
            if isinstance(value, float):
                value = round(value, 3)
            cell.text = str(value)

    # Adjust the width of columns
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)


def underline_text(document, text, underline_style=WD_UNDERLINE.SINGLE):
    """
    Add underlined text to a Word document.

    Args:
    - document: The Document object where the underlined text will be added.
    - text: The text to be underlined.
    - underline_style: The style of underline. Default is WD_UNDERLINE.SINGLE.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.underline = underline_style

def add_right_aligned_bold_text(document, text, font_size=12):
    """
    Add right-aligned bold text to a Word document.

    Args:
    - document: The Document object where the text will be added.
    - text: The text to be added.
    - font_size: The font size of the text. Default is 12.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font = run.font
    font.size = Pt(font_size)
    
def int_to_roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num

