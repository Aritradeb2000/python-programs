from docx import Document
import textwrap
import datetime
import os
import errno

#Use for time limit
expiration_date = datetime.datetime(2024, 7, 15)  
def check_expiration():
    current_date = datetime.datetime.now()
    if current_date > expiration_date:
        print("This program has expired.")
        exit()

check_expiration()


import input_sheet
import time_period_cal
import abut_foundation
import sls_check_seismic
import sls_check_nonseismic

user_home = os.path.expanduser("~")

output_folder = os.path.join("D:", "Python", "Abutment Foundation(Pile)")

try:
    os.makedirs(output_folder)
except OSError as e:
    if e.errno != errno.EEXIST:
        print(f"Error creating directory: {e}")
        output_folder = os.path.join(user_home, 'Python', 'Abutment Foundation(Pile)')
    else:
        pass
filename = input("\n\nPlease Specify Output file name(without extension): ")
output_file_path = os.path.join(output_folder, f'{filename}.docx')
   

document = Document()




print(f'Document will be saved at: {output_file_path}')


document = Document()
# document.add_paragraph(f"This is the design of an Abutment wall. The wall is resting on well of suitable inside and outside diameter.\n\nAs a first step,the formation level is defined and there after the rail levels are calculated as per the standard IRC spec. The weight of girder(which is an open wen girder) is taken up from IRC spec. Live loads ,Breaking loads and Tractive Forces are also taken into account. Weight of pedestal, bearing etc are also taken care off.\n\nSeismic analysis is as per latest IS code for the pear.Seismic forces up to the scour level. Soil investigation is carried out at site in basic design parameters are given here.However  the report was submitted and got approved by thr department before proceeding with the design. The summarised soil test results are also  incorporated.\n\nFor the convinence all the input parameters are placed to the beging dor future refference during compution compution will automatically seek and use the parameter. Intermittent use of input parameters are avoided for accuracy and speed.\n\nDesign of well foundation was very critical to bending moment due to high seismic forces. However due to huge dimention and proper reinforcement the well did not crack and stresses remain positive.\n\nPile formations were disregurded because of seismic forces and presense of moderate sand.\n\n")




while(True):
    print("Enter 0 to exit the program.")
    
    print("Enter 1 to print input Sheet")

    print("Enter 2 to print abutment foundation sheet")

    print("Enter 3 to print SLS_check(Seismic) sheet")

    print("Enter 4 to print SLS_check(Nonseismic) sheet")
    

    try:
        print('\n\n')
        user_input = int(input("Please enter the design request(0-4): "))
        print("\n\n")

        if user_input < 0 or user_input >4:
            print("Please enter design value.")
            continue
    except:
        print("Please Enter Correct input.")
        continue
    if(user_input == 0):
        print("Exiting Program")
        break
    elif(user_input == 1):
        print("Printing input Sheet")
        input_sheet.word_print(document,output_file_path)
        time_period_cal.word_print(document,output_file_path)
    elif(user_input == 2):
        print("printing abutment and foundation sheet")
        abut_foundation.word_print(document, output_file_path)
    elif(user_input == 3):
        print("printing SLS_check(Seismic) sheet")
        sls_check_seismic.word_print(document, output_file_path)
    elif(user_input == 4):
        print("printing SLS_check(Nonseismic) sheet")
        sls_check_nonseismic.word_print(document, output_file_path)
    
        
    else:
        print("Please Enter valid Input")
           
    
    
    
    
document.save(output_file_path)

